import sys
import os

# this is needed to import classes from other modules
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../')))

import requests
import magic # magic requires lib magic to be installed on the system. If running on mac and an error occurs, run `brew install libmagic`
import json
import fitz
import logging
import copy
from io import BytesIO
import services.database.batch_service as batch_service
import services.database.job_service as job_service
from flask import Flask, request, jsonify
from flask_cors import CORS
from models.batch import Batch
from api.auth import Auth
from api.pipeline import Pipeline
from services.database.database import get_db, safe_db_operation
from shared.vectorflow_request import VectorflowRequest
from docx import Document
from urllib.parse import urlparse
from pathlib import Path
from services.minio.minio_service import create_minio_client
from api.posthog import send_telemetry
from datetime import datetime
from api.validators import RequestValidator, Validations

auth = Auth()
pipeline = Pipeline()
app = Flask(__name__)
CORS(app)

logging.basicConfig(filename='./api-log.txt', level=logging.INFO)
logging.basicConfig(filename='./api-errors.txt', level=logging.ERROR)


@app.route("/embed", methods=['POST'])
def embed():
    vectorflow_request = VectorflowRequest._from_flask_request(request)
    if invalid := RequestValidator(request, auth).validate([Validations.CRED, 
                                                            Validations.METADATA, 
                                                            Validations.EMBEDDING_TYPE, 
                                                            Validations.WEBHOOK, 
                                                            Validations.SOURCE_DATA]):
        return RequestValidator.dispatch_on_invalid(invalid, jsonify)

    file = request.files['SourceData']
    
    # TODO: Remove this once the application is reworked to support large files
    # Get the file size - Go to the end of the file, get the current position, and reset the file to the beginning
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)

    if file_size > 25 * 1024 * 1024:
        return jsonify({'error': 'File is too large. The /embed endpoint currently only supports 25 MB files or less. Please use /jobs for streaming large files or multiple files.'}), 413
    
    # empty filename means no file was selected
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and is_valid_file_type(file):
        job = safe_db_operation(job_service.create_job, vectorflow_request, file.filename)
        batch_count = process_file(file, vectorflow_request, job.id)

        vectorflow_request_copy = copy.deepcopy(vectorflow_request)
        send_telemetry("SINGLE_FILE_UPLOAD_SUCCESS", vectorflow_request_copy)

        logging.info(f"{datetime.now()} Successfully created job {job.id} for file {file.filename}")
        return jsonify({'message': f"Successfully added {batch_count} batches to the queue", 'JobID': job.id}), 200
    else:
        return jsonify({'error': 'Uploaded file is not a TXT, PDF, Markdown or DOCX file'}), 400

@app.route('/jobs', methods=['POST'])
def create_jobs():
    vectorflow_request = VectorflowRequest._from_flask_request(request)
    if invalid := RequestValidator(request, auth).validate([Validations.CRED, 
                                                            Validations.METADATA, 
                                                            Validations.EMBEDDING_TYPE, 
                                                            Validations.WEBHOOK, 
                                                            Validations.HAS_FILES]):
        return RequestValidator.dispatch_on_invalid(invalid, jsonify)

    files = request.files.getlist('file')
    successfully_uploaded_files = dict()
    failed_uploads = []
    empty_files_count = 0
    duplicate_files_count = 0

    for file in files:
        # Check if a file is empty (no filename)
        if file.filename == '':
            empty_files_count += 1
            continue
        
        if not is_valid_file_type(file):
            failed_uploads.append(file.filename)
            continue

        if file.filename in successfully_uploaded_files:
            duplicate_files_count += 1
            continue

        temporary_storage_location = os.getenv('API_STORAGE_DIRECTORY')
        file_path = os.path.join(temporary_storage_location, file.filename)
        with open(file_path, 'wb') as f:
            chunk_size = 65536  # 64 KB
            while True:
                chunk = file.stream.read(chunk_size)
                if len(chunk) == 0:
                    break
                f.write(chunk)
        try:
            result = upload_to_minio(file_path, file.filename)
            os.remove(file_path)

            if not result:
                failed_uploads.append(file.filename)
                continue
            
            job = safe_db_operation(job_service.create_job, vectorflow_request, file.filename)
            if not job:
                remove_from_minio(file.filename)
                continue
       
            data = (job.id, file.filename, vectorflow_request.serialize())
            json_data = json.dumps(data)

            pipeline.connect(queue=os.getenv('EXTRACTION_QUEUE'))
            pipeline.add_to_queue(json_data, queue=os.getenv('EXTRACTION_QUEUE'))
            pipeline.disconnect()

            successfully_uploaded_files[file.filename] = job.id
            logging.info(f"{datetime.now()} Successfully created job {job.id} for file {file.filename}")
        except Exception as e:
            print(f"Error uploading file {file.filename} to min.io, creating job or passing vectorflow request to message broker. \nError: {e}\n\n")
            failed_uploads.append(file.filename)       

    send_telemetry("MULTI_FILE_UPLOAD_SUCCESS", vectorflow_request)
    return jsonify({'message': 'Files processed', 
                    'successful_uploads': successfully_uploaded_files,
                    'failed_uploads': failed_uploads,
                    'empty_files_count': empty_files_count,
                    'duplicate_files_count': duplicate_files_count}), 200

@app.route('/jobs/<int:job_id>/status', methods=['GET'])
def get_job_status(job_id):
    vectorflow_key = request.headers.get('Authorization')
    if not vectorflow_key or not auth.validate_credentials(vectorflow_key):
        return jsonify({'error': 'Invalid credentials'}), 401
    
    job = safe_db_operation(job_service.get_job, job_id)
    if job:
        return jsonify({'JobStatus': job.job_status.value}), 200
    else:
        return jsonify({'error': "Job not found"}), 404
        
@app.route('/jobs/status', methods=['POST'])
def get_job_statuses():
    vectorflow_key = request.headers.get('Authorization')
    if not vectorflow_key or not auth.validate_credentials(vectorflow_key):
        return jsonify({'error': 'Invalid credentials'}), 401
    
    if not hasattr(request, 'json') or not request.json:
        return jsonify({'error': 'Missing JSON body'}), 400
    
    job_ids = request.json.get('JobIDs')
    if not job_ids:
        return jsonify({'error': 'Missing JobIDs field'}), 400
    
    jobs = safe_db_operation(job_service.get_jobs, job_ids)
    if jobs:
        return jsonify({'Jobs': [{'JobID': job.id, 'JobStatus': job.job_status.value} for job in jobs]}), 200
    else:
        return jsonify({'error': "Jobs not found"}), 404   
    
@app.route("/s3", methods=['POST'])
def s3_presigned_url():
    vectorflow_request = VectorflowRequest._from_flask_request(request)
    if invalid := RequestValidator.validate([Validations.CRED, 
                                             Validations.WEBHOOK, 
                                             Validations.EMBEDDING_TYPE, 
                                             Validations.PRE_SIGNED]):
        return RequestValidator.dispatch_on_invalid(invalid, jsonify)

    pre_signed_url = request.form.get('PreSignedURL')
    
    response = requests.get(pre_signed_url)
    file_name = get_s3_file_name(pre_signed_url)

    if response.status_code == 200:
        file_magic = magic.Magic(mime=True)
        mime_type = file_magic.from_buffer(response.content)

        if mime_type == 'text/plain':
            file_content = response.text
            job = safe_db_operation(job_service.create_job, vectorflow_request, file_name)
            return jsonify({'message': f"Successfully added {batch_count} batches to the queue", 'JobID': job.id}), 200
        
        elif mime_type == 'application/pdf':
            pdf_data = BytesIO(response.content)
            with fitz.open(stream=pdf_data, filetype='pdf') as doc:
                file_content = ""
                for page in doc:
                    file_content += page.get_text()

            job = safe_db_operation(job_service.create_job, vectorflow_request, file_name)
            batch_count = create_batches(file_content, job.id, vectorflow_request)
            return jsonify({'message': f"Successfully added {batch_count} batches to the queue", 'JobID': job.id}), 200
        
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            docx_data = BytesIO(response.content)
            doc = Document(docx_data)
            file_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            job = safe_db_operation(job_service.create_job, vectorflow_request, file_name)
            batch_count = create_batches(file_content, job.id, vectorflow_request)
            return jsonify({'message': f"Successfully added {batch_count} batches to the queue", 'JobID': job.id}), 200
        
        else:
            return jsonify({'error': 'Uploaded file is not a TXT, PDF, HTML or DOCX file'}), 400
    else:
        print('Failed to download file:', response.status_code, response.reason)

def process_file(file, vectorflow_request, job_id):
    if file.filename.endswith('.txt'):
        file_content = file.read().decode('utf-8')
    
    elif file.filename.endswith('.docx'):
        doc = Document(file)
        file_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    elif file.filename.endswith('.md'):
        temp_file_path = Path('./temp_file.md')
        file.save(temp_file_path)
            
        with open(temp_file_path, 'r') as file:
            file_content = file.read()
        
        temp_file_path.unlink()
    
    elif file.filename.endswith('.html'):
        content = file.read().decode('utf-8')
        file_content = repr(content)

    else:
        pdf_data = BytesIO(file.read())
        with fitz.open(stream=pdf_data, filetype='pdf') as doc:
            file_content = ""
            for page in doc:
                file_content += page.get_text()

    batch_count = create_batches(file_content, job_id, vectorflow_request)
    return batch_count

def create_batches(file_content, job_id, vectorflow_request_original):
    vectorflow_request = copy.deepcopy(vectorflow_request_original)
    chunks = [chunk for chunk in split_file(file_content, vectorflow_request.lines_per_batch)]
    
    batches = [Batch(job_id=job_id, embeddings_metadata=vectorflow_request.embeddings_metadata, vector_db_metadata=vectorflow_request.vector_db_metadata) for _ in chunks]
    batches = safe_db_operation(batch_service.create_batches, batches)

    job = safe_db_operation(job_service.update_job_total_batches, job_id, len(batches))

    for batch, chunk in zip(batches, chunks):
        data = (batch.id, chunk, vectorflow_request.vector_db_key, vectorflow_request.embedding_api_key)
        json_data = json.dumps(data)

        pipeline.connect(queue=os.getenv('EMBEDDING_QUEUE'))
        pipeline.add_to_queue(json_data, queue=os.getenv('EMBEDDING_QUEUE'))
        pipeline.disconnect()

    return job.total_batches if job else None
    
def split_file(file_content, lines_per_chunk=1000):
    lines = file_content.splitlines()
    for i in range(0, len(lines), lines_per_chunk):
        yield lines[i:i+lines_per_chunk]
    
def get_s3_file_name(pre_signed_url):
    parsed_url = urlparse(pre_signed_url)
    path_parts = parsed_url.path.lstrip('/').split('/')

    # For the file name and not the full path:
    file_name = path_parts[-1]
    return file_name

def is_valid_file_type(file):
    supported_types = ['.txt', '.docx', '.pdf', '.md', '.html']
    for type in supported_types:
        if file.filename.endswith(type):
            return True
        
    # Try to detect .txt files by content
    try:
        # Read a portion of the file
        file_content = file.stream.read(1024)
        file.stream.seek(0)  # Reset the file stream position for later use

        # Attempt to decode the content as utf-8 text
        file_content.decode('utf-8')

        # If we were able to successfully decode the file as utf-8 text, it's likely a text file
        file.filename += '.txt'
        return True  # Successful decoding implies it's a text file
    except UnicodeDecodeError:
        return False  # Failed to decode, likely not a text file

def remove_from_minio(filename):
    client = create_minio_client()
    client.remove_object(os.getenv("MINIO_BUCKET"), filename)

def upload_to_minio(file_path, filename):
    client = create_minio_client()

    file_size = os.path.getsize(file_path)

    # Wrap the generator with our StreamWrapper
    stream = StreamWrapper(lambda: file_data_generator(file_path))

    result = client.put_object(
        os.getenv("MINIO_BUCKET"), filename, stream, file_size
    )
    return result

# generator used to stream
def file_data_generator(file_path, chunk_size=65536):  # 64KB
    with open(file_path, 'rb') as file:
        while True:
            data = file.read(chunk_size)
            if not data:
                break
            yield data

class StreamWrapper:
    def __init__(self, generator_func):
        self.generator = generator_func()

    def read(self, *args):
        return next(self.generator, b'')

if __name__ == '__main__':
   app.run(host='0.0.0.0', debug=True)